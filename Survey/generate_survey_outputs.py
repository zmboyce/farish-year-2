#!/usr/bin/env python3
"""
Farish Street Community Survey – Year 2 outputs (+ Year 1 vs Year 2 comparisons).

Produces:
  Survey/Outputs/          – bar‑chart PNGs matching the Year 1 PDF style
  Survey/survey_dashboard.html  – interactive comparison dashboard
  Survey/Outputs/Farish_Survey_Year2_Report.docx  – Word doc in PDF sequence

Run from repo root:
    python Survey/generate_survey_outputs.py
"""

from __future__ import annotations

import datetime as dt
import html
import json
import re
import textwrap
from dataclasses import dataclass, field
from pathlib import Path

import matplotlib.font_manager as fm
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np
import pandas as pd

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "Outputs"
FONTS = ROOT.parent / "fonts"
Y1_PATH = ROOT / "Farish St Results Summary - Year 1.xlsx"
Y2_PATH = ROOT / "Farish St Results Summary - Year 2.xlsx"

# ── Exact colors from Excel chart XMLs ───────────────────────────────────────
COL_DATE1  = "#6B9CCC"   # June 22 / Year 1 first series
COL_DATE2  = "#197A80"   # August 25 / Year 1 second series
COL_ALL    = "#6C72B6"   # All festival dates / Year 1 combined
COL_Y2     = "#197A80"   # Year 2 aggregate (single series – teal, same as August25)
COL_Y2_COMP = "#C0504D"  # used only in head-to-head Year 1 vs Year 2 bar

DATE_LABELS_Y1 = ("June 22", "August 25", "All festival dates")

# ── Titles exactly as they appear in the Year 1 PDF ──────────────────────────
Q_TITLES: dict[int, str] = {
    1:  "1. What is your age",
    2:  "2. How long have you lived in Jackson?",
    3:  "3. What is your race/ethnicity?",
    4:  "4. What is your yearly household income?",
    5:  "5. Did you participate in any outdoor events on\nFarish Street last summer?",
    6:  "6. Were you more, less, or equally comfortable with\nthe heat and humidity on Farish Street last year?",
    7:  "7. Considering the combination of heat and\nhumidity, how does it feel to be outside on\nFarish Street today?",
    8:  None,          # matrix – subquestion used as title
    9:  "9. Do you ever decide to stay home from outdoor\nevents because it is too hot outside?",
    11: "11. Are you more likely to spend time outside in the\nsummer if you have a green space to go to, like a\npark, garden, or a grassy field?",
    13: "13. Do you prefer parks and green spaces in\nJackson with or without trees?",
    14: "14. Do you have any concerns related to trees in\npublic parks?",
    15: "15. Do you think it would be more enjoyable to spend\ntime on Farish Street if there was a park or green\nspace to visit?",
}

SECTION_HEADINGS: dict[str, str] = {
    "SocioDem":        "Socio-Demographics",
    "Comfortability":  "Comfort & Heat Impacts",
    "Decision Making": "Decision Making",
    "Green Spaces":    "Green Space Perceptions",
}

# Q number inferred from question text
def _q_num(qtext: str) -> int | None:
    m = re.match(r"(\d+)\.", qtext.strip())
    return int(m.group(1)) if m else None


# ── Font setup ────────────────────────────────────────────────────────────────
def setup_fonts() -> str:
    for p in FONTS.glob("*.otf"):
        fm.fontManager.addfont(str(p))
    available = {f.name for f in fm.fontManager.ttflist}
    for candidate in [
        "Avenir Next LT Pro",
        "AvenirNextLTPro-Cn",
        "Avenir Next LT Pro Condensed",
    ]:
        if candidate in available:
            return candidate
    avenir = [n for n in available if "avenir" in n.lower()]
    return avenir[0] if avenir else "DejaVu Sans"


def apply_style(font: str) -> None:
    plt.rcParams.update({
        "font.family": font,
        "font.size": 10,
        "figure.dpi": 150,
        "axes.spines.top": False,
        "axes.spines.right": False,
        "axes.spines.left": False,
        "axes.spines.bottom": True,
        "axes.linewidth": 0.7,
        "axes.edgecolor": "#888888",
        "axes.grid": True,
        "axes.grid.axis": "y",
        "axes.axisbelow": True,   # grid lines drawn behind bars
        "grid.color": "#e0e0e0",
        "grid.linewidth": 0.6,
        "xtick.major.size": 0,
        "ytick.major.size": 0,
        "xtick.color": "#555555",
        "ytick.color": "#555555",
        "axes.labelcolor": "#333333",
        "text.color": "#222222",
    })


# ── Parsing helpers ───────────────────────────────────────────────────────────
def _slug(s: str, maxlen: int = 48) -> str:
    s = re.sub(r"[^\w\s-]", "", str(s).lower())[:maxlen]
    return re.sub(r"[-\s]+", "_", s).strip("_") or "q"


def is_question_cell(c0) -> bool:
    if pd.isna(c0):
        return False
    s = str(c0).strip()
    return bool(re.match(r"^\d+\.", s)) or s.lower().startswith("multi-matrix")


def is_date_label(c0) -> str | bool:
    if pd.isna(c0):
        return False
    if isinstance(c0, (pd.Timestamp, dt.datetime)):
        return "date"
    s = str(c0).strip().lower()
    if "all festival" in s or "all event" in s:
        return "all"
    return False


def is_percent_row(vals: list[float]) -> bool:
    if not vals or any(np.isnan(x) for x in vals):
        return False
    return all(0 <= x <= 1.05 for x in vals)


def is_duplicate_scale_header(vals: list[float]) -> bool:
    if not vals or any(np.isnan(x) for x in vals):
        return False
    try:
        return [int(float(x)) for x in vals] == list(range(1, len(vals) + 1))
    except (TypeError, ValueError):
        return False


def is_legend_row(row: pd.Series, opt_cols: list[int]) -> bool:
    c1 = row.iloc[1] if len(row) > 1 else np.nan
    try:
        v = float(c1)
    except (TypeError, ValueError):
        return False
    if not v.is_integer() or not (1 <= int(v) <= 20):
        return False
    return all(pd.isna(row.iloc[j]) for j in opt_cols[1:] if j < len(row))


def option_cells_all_na(row: pd.Series, opt_cols: list[int]) -> bool:
    return all(pd.isna(row.iloc[j]) for j in opt_cols if j < len(row))


def _read_counts(row: pd.Series, opt_cols: list[int]) -> list[float] | None:
    vals: list[float] = []
    for j in opt_cols:
        if j >= len(row):
            return None
        v = row.iloc[j]
        if pd.isna(v):
            vals.append(np.nan)
        elif isinstance(v, (int, float, np.integer, np.floating)):
            vals.append(float(v))
        else:
            return None
    if all(np.isnan(x) for x in vals):
        return None
    return vals


def peel_matrix_trailer(qtext: str) -> tuple[str, str | None]:
    lines = [ln.strip() for ln in str(qtext).split("\n") if ln.strip()]
    if len(lines) >= 2:
        last = lines[-1]
        if (
            len(last) < 70
            and "?" not in last
            and not re.match(r"^\d+\.", last)
            and "multi-matrix" not in last.lower()
        ):
            return "\n".join(lines[:-1]), last
    return str(qtext).strip(), None


def is_matrix_question(q: str) -> bool:
    q = q.lower()
    return "multi-matrix" in q


def extract_option_columns(header: pd.Series, qtext: str) -> list[int]:
    cols: list[int] = []
    qshort = str(qtext)[:40].lower()
    for j in range(1, min(len(header), 16)):
        v = header.iloc[j]
        if pd.isna(v):
            if cols:
                break
            continue
        if isinstance(v, str):
            vs = v.strip().lower()
            if len(vs) > 12 and qshort.startswith(vs[:12]):
                break
        cols.append(j)
    return cols


def normalize_question_key(text: str) -> str:
    t = str(text).lower()
    t = re.sub(r"\s+", " ", t)
    t = t.replace("hispanic/ latine", "hispanic/ latino")
    t = t.replace("comfotable", "comfortable")
    return t[:220]


# ── Data structures ───────────────────────────────────────────────────────────
@dataclass
class SeriesRow:
    label: str
    counts: list[float]
    percents: list[float] | None = None


@dataclass
class QuestionBlock:
    sheet: str
    question: str
    subquestion: str | None
    option_labels: list[str]
    series: list[SeriesRow] = field(default_factory=list)
    year_tag: str = ""

    def q_num(self) -> int | None:
        return _q_num(self.question)

    def key(self) -> str:
        base = normalize_question_key(self.question)
        if self.subquestion:
            base += "||" + normalize_question_key(self.subquestion).replace(" ", "_")
        return base

    def plot_title(self) -> str:
        qn = self.q_num()
        if is_matrix_question(self.question) and self.subquestion:
            return self.subquestion
        if qn and qn in Q_TITLES and Q_TITLES[qn]:
            return Q_TITLES[qn]
        return re.sub(r"^Multi-matrix question\n", "", self.question).strip()


# ── Parser ────────────────────────────────────────────────────────────────────
def parse_sheet(df: pd.DataFrame, sheet: str, year_tag: str) -> list[QuestionBlock]:
    blocks: list[QuestionBlock] = []
    i = 0
    n = len(df)

    while i < n:
        c0 = df.iloc[i, 0]
        if not is_question_cell(c0):
            i += 1
            continue
        qtext = str(c0).strip()
        if qtext.lower().startswith("clarify "):
            i += 1
            continue
        header = df.iloc[i]
        opt_cols = extract_option_columns(header, qtext)
        if not opt_cols:
            i += 1
            continue
        raw_opts = [str(header.iloc[j]).strip() for j in opt_cols]
        q_base, matrix_first_sub = peel_matrix_trailer(qtext)
        block = QuestionBlock(
            sheet=sheet,
            question=q_base,
            subquestion=matrix_first_sub,
            option_labels=raw_opts,
            series=[],
            year_tag=year_tag,
        )
        i += 1
        likert_map: dict[int, str] = {}

        def apply_likert(b: QuestionBlock) -> None:
            if not likert_map:
                return
            new_opts = []
            for x in b.option_labels:
                lab = str(x).strip()
                code = None
                if lab.isdigit():
                    code = int(lab)
                elif re.match(r"^\d+\.0+$", lab):
                    code = int(float(lab))
                new_opts.append(likert_map.get(code, lab) if code else lab)
            b.option_labels = new_opts

        while i < n:
            row = df.iloc[i]
            c = row.iloc[0]

            if is_question_cell(c):
                break

            c_str = None if pd.isna(c) else str(c).strip()

            if isinstance(c_str, str):
                if c_str.lower().startswith("clarify "):
                    i += 1
                    continue
                if c_str.startswith("Question ") and "Write" in c_str:
                    i += 1
                    continue
                if "of these," in c_str.lower() and "were" in c_str.lower():
                    i += 1
                    continue
                if c_str.lower().startswith("just numbers"):
                    i += 1
                    continue

            all_na = option_cells_all_na(row, opt_cols)

            if (
                c_str
                and len(c_str) < 90
                and "?" not in c_str
                and not is_date_label(c)
                and not re.match(r"^\d{4}-\d{2}-\d{2}", c_str)
            ):
                vals_try = _read_counts(row, opt_cols)
                if all_na and c_str and not is_legend_row(row, opt_cols):
                    if is_matrix_question(block.question) and c_str:
                        if block.series:
                            apply_likert(block)
                            blocks.append(block)
                        block = QuestionBlock(
                            sheet=sheet,
                            question=q_base,
                            subquestion=c_str,
                            option_labels=list(raw_opts),
                            series=[],
                            year_tag=year_tag,
                        )
                    i += 1
                    continue
                if vals_try and not is_percent_row(vals_try):
                    if is_legend_row(row, opt_cols):
                        try:
                            code = int(float(row.iloc[1]))
                            likert_map[code] = c_str
                        except (TypeError, ValueError):
                            pass
                        i += 1
                        continue
                    if is_matrix_question(block.question):
                        if block.subquestion is not None and c_str != block.subquestion:
                            if block.series:
                                apply_likert(block)
                                blocks.append(block)
                            block = QuestionBlock(
                                sheet=sheet,
                                question=q_base,
                                subquestion=c_str,
                                option_labels=list(raw_opts),
                                series=[],
                                year_tag=year_tag,
                            )
                        elif block.subquestion is None:
                            block.subquestion = c_str
                        block.series.append(SeriesRow(label=year_tag + " (all respondents)", counts=vals_try))
                        i += 1
                        continue
                    block.series.append(SeriesRow(label=c_str, counts=vals_try))
                    i += 1
                    continue

            dlab = is_date_label(c)
            if dlab:
                vals = _read_counts(row, opt_cols)
                if vals and not is_percent_row(vals):
                    if dlab == "date":
                        lbl = pd.Timestamp(c).strftime("%b %-d")
                    else:
                        lbl = "All festival dates"
                    block.series.append(SeriesRow(label=lbl, counts=vals))
                i += 1
                continue

            if pd.isna(c) or not c_str or c_str == "nan":
                vals = _read_counts(row, opt_cols)
                if vals and is_duplicate_scale_header(vals):
                    i += 1
                    continue
                if vals and not is_percent_row(vals):
                    block.series.append(SeriesRow(label=year_tag + " (all respondents)", counts=vals))
                    i += 1
                    continue
                if vals and is_percent_row(vals):
                    if block.series:
                        block.series[-1].percents = vals
                    i += 1
                    continue
                i += 1
                continue

            i += 1

        apply_likert(block)
        if block.series:
            blocks.append(block)
    return blocks


def clean_block(block: QuestionBlock) -> QuestionBlock | None:
    if "10. If YES" in block.question and "temperature" in block.question.lower():
        return None
    if not block.option_labels:
        return None
    seen: set[tuple] = set()
    series2: list[SeriesRow] = []
    for s in block.series:
        key = (s.label, tuple(s.counts))
        if key not in seen:
            seen.add(key)
            series2.append(s)
    block.series = series2
    return block if series2 else None


def load_all_blocks(path: Path, year: str) -> dict[str, list[QuestionBlock]]:
    xl = pd.ExcelFile(path)
    out: dict[str, list[QuestionBlock]] = {}
    for sh in xl.sheet_names:
        df = pd.read_excel(path, sheet_name=sh, header=None)
        out[sh] = [b for b in (clean_block(b) for b in parse_sheet(df, sh, year)) if b]
    return out


# ── Count / percent helpers ───────────────────────────────────────────────────
def counts_to_percent(counts: list[float]) -> list[float]:
    arr = np.array(counts, dtype=float)
    s = np.nansum(arr)
    return list(np.nan_to_num(arr / s, nan=0.0)) if s > 0 else [0.0] * len(counts)


def _pct_list(s: SeriesRow, n_opts: int) -> list[float]:
    if s.percents and len(s.percents) >= n_opts and not any(np.isnan(x) for x in s.percents[:n_opts]):
        return [x * 100 if x <= 1 else x for x in s.percents[:n_opts]]
    return [c * 100 for c in counts_to_percent(s.counts[:n_opts])]


# ── Plot helpers ──────────────────────────────────────────────────────────────
def _series_colors(labels: list[str]) -> list[str]:
    """Match series labels to the Year 1 Excel bar palette."""
    result: list[str] = []
    for lab in labels:
        ll = lab.lower()
        if "june" in ll or "jun" in ll:
            result.append(COL_DATE1)
        elif "aug" in ll:
            result.append(COL_DATE2)
        elif "all festival" in ll or "all event" in ll:
            result.append(COL_ALL)
        elif "year 2" in ll or "y2" in ll:
            result.append(COL_Y2)
        else:
            # cycle through palette
            palette = [COL_DATE1, COL_DATE2, COL_ALL]
            result.append(palette[len(result) % 3])
    return result


def plot_block(
    block: QuestionBlock,
    out_path: Path,
    year_label: str = "Year 2",
) -> None:
    """Grouped vertical bar chart matching Year 1 PDF style."""
    opts = block.option_labels
    series = block.series
    if not opts or not series:
        return

    n_o = len(opts)
    labels = [s.label for s in series]
    pcts = [_pct_list(s, n_o) for s in series]
    n_s = len(pcts)

    colors = _series_colors(labels)
    width = min(0.80 / max(n_s, 1), 0.25)
    fig_w = max(7.5, 0.5 * n_o * n_s + 2.5)
    fig, ax = plt.subplots(figsize=(fig_w, 4.8))

    x = np.arange(n_o)
    for si, (lab, pc) in enumerate(zip(labels, pcts)):
        offset = width * (si - (n_s - 1) / 2)
        ax.bar(
            x + offset, pc[:n_o], width,
            label=lab,
            color=colors[si],
            edgecolor="white",
            linewidth=0.4,
        )

    ax.set_ylabel("Percent of respondents", fontsize=9)
    ymax = max(max(p[:n_o]) for p in pcts) if pcts else 100
    ax.set_ylim(0, min(100, ymax * 1.18 + 3))
    ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda v, _: f"{v:.0f}%"))
    ax.set_xticks(x)
    wrapped = [textwrap.fill(o.replace("/", "/ "), 14) for o in opts]
    ax.set_xticklabels(wrapped, fontsize=8.5)
    ax.tick_params(bottom=False)
    ax.spines["bottom"].set_visible(False)
    ax.set_xlabel("Response", fontsize=9, color="#555555", labelpad=6)

    title = block.plot_title()
    ax.set_title(title, fontsize=13, fontweight="bold", pad=10, loc="center")

    if n_s > 1:
        handles = [mpatches.Patch(color=colors[si], label=labels[si]) for si in range(n_s)]
        ax.legend(handles=handles, loc="upper right", fontsize=8, framealpha=0.9, edgecolor="#dddddd")

    fig.tight_layout(pad=1.2)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(out_path, dpi=150, bbox_inches="tight")
    plt.close(fig)


def plot_y1_y2_comparison(b1: QuestionBlock, b2: QuestionBlock, out_path: Path) -> None:
    """Paired bars: Year 1 (all festival dates) vs Year 2 aggregate."""
    s1 = next((s for s in b1.series if "All festival" in s.label), b1.series[-1] if b1.series else None)
    s2 = next((s for s in b2.series if "all respondents" in s.label.lower()), b2.series[0] if b2.series else None)
    if not s1 or not s2:
        return
    opts1, opts2 = b1.option_labels, b2.option_labels
    if len(opts1) != len(opts2):
        return
    n = min(len(opts1), len(s1.counts), len(s2.counts))
    if n == 0:
        return
    opts = opts1[:n]
    p1 = [x * 100 for x in counts_to_percent(s1.counts[:n])]
    p2 = [x * 100 for x in counts_to_percent(s2.counts[:n])]

    w = 0.36
    x = np.arange(n)
    fig, ax = plt.subplots(figsize=(max(7.5, 0.55 * n + 3.5), 4.8))
    ax.bar(x - w / 2, p1, w, label="Year 1 (all festival dates)", color=COL_ALL, edgecolor="white")
    ax.bar(x + w / 2, p2, w, label="Year 2 (all respondents)", color=COL_Y2_COMP, edgecolor="white")
    ax.set_ylabel("Percent of respondents", fontsize=9)
    ax.set_ylim(0, 100)
    ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda v, _: f"{v:.0f}%"))
    ax.set_xticks(x)
    ax.set_xticklabels([textwrap.fill(o.replace("/", "/ "), 16) for o in opts], fontsize=8.5)
    ax.tick_params(bottom=False)
    ax.spines["bottom"].set_visible(False)
    ax.set_xlabel("Response", fontsize=9, color="#555555", labelpad=6)
    title = b1.plot_title()
    if b1.subquestion and not is_matrix_question(b1.question):
        title += f" — {b1.subquestion}"
    ax.set_title(title, fontsize=13, fontweight="bold", pad=10, loc="center")
    ax.legend(fontsize=8.5, framealpha=0.9, edgecolor="#dddddd")
    ax.grid(axis="y", linestyle=":", alpha=0.5)
    ax.spines["left"].set_visible(False)
    fig.tight_layout(pad=1.2)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(out_path, dpi=150, bbox_inches="tight")
    plt.close(fig)


# ── Data table builder ────────────────────────────────────────────────────────
def build_table_df(block: QuestionBlock) -> pd.DataFrame:
    """Count table: rows = series / dates, columns = options."""
    rows = []
    for s in block.series:
        row = {"Date": s.label}
        for j, opt in enumerate(block.option_labels):
            row[opt] = int(round(s.counts[j])) if j < len(s.counts) and not np.isnan(s.counts[j]) else 0
        rows.append(row)
    return pd.DataFrame(rows).set_index("Date")


# ── Comparison index ──────────────────────────────────────────────────────────
def build_comparison_index(
    y1: dict[str, list[QuestionBlock]],
    y2: dict[str, list[QuestionBlock]],
) -> list[tuple[str, str, QuestionBlock, QuestionBlock]]:
    comp: list[tuple] = []
    for sheet in y1:
        if sheet not in y2:
            continue
        m2 = {b.key(): b for b in y2[sheet]}
        for b1 in y1[sheet]:
            k = b1.key()
            if k in m2:
                comp.append((sheet, k, b1, m2[k]))
    return comp


def build_change_summaries(
    comp: list[tuple[str, str, QuestionBlock, QuestionBlock]],
) -> list[dict]:
    rows: list[dict] = []
    for sheet, _key, b1, b2 in comp:
        if "10. if yes" in b1.question.lower():
            continue
        s1 = next((s for s in b1.series if "All festival" in s.label), b1.series[-1] if b1.series else None)
        s2 = next((s for s in b2.series if "all respondents" in s.label.lower()), b2.series[0] if b2.series else None)
        if not s1 or not s2:
            continue
        p1 = counts_to_percent(s1.counts)
        p2 = counts_to_percent(s2.counts)
        n = min(len(p1), len(p2), len(b1.option_labels), len(b2.option_labels))
        if n < 2:
            continue
        rows.append({
            "sheet": sheet,
            "title": (b1.subquestion + " · " if b1.subquestion else "") + b1.question.replace("\n", " ")[:80],
            "options": b1.option_labels[:n],
            "year1_pct": [round(p1[i] * 100, 1) for i in range(n)],
            "year2_pct": [round(p2[i] * 100, 1) for i in range(n)],
            "delta_pct_points": [round(p2[i] * 100 - p1[i] * 100, 1) for i in range(n)],
        })
    return rows


# ── Word document ─────────────────────────────────────────────────────────────
def build_word_doc(
    y2: dict[str, list[QuestionBlock]],
    png_paths: dict[str, Path],
    font_name: str,
) -> None:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    # Default paragraph font
    style = doc.styles["Normal"]
    style.font.name = font_name if font_name != "DejaVu Sans" else "Calibri"
    style.font.size = Pt(11)

    def add_heading(text: str, level: int = 1) -> None:
        p = doc.add_heading(text, level=level)
        run = p.runs[0] if p.runs else p.add_run(text)
        run.font.name = font_name if font_name != "DejaVu Sans" else "Calibri"
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        if level == 1:
            run.font.size = Pt(14)
        else:
            run.font.size = Pt(11)

    def add_body(text: str) -> None:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)

    def add_section_divider(heading: str) -> None:
        doc.add_page_break()
        add_heading(heading, level=1)

    def add_table(df: pd.DataFrame) -> None:
        if df.empty:
            return
        n_cols = len(df.columns) + 1
        tbl = doc.add_table(rows=1 + len(df), cols=n_cols)
        tbl.style = "Light Shading Accent 1"
        # Header row
        hdr = tbl.rows[0].cells
        hdr[0].text = "Date"
        for j, col in enumerate(df.columns):
            hdr[j + 1].text = str(col)
        for r_idx, (idx, row) in enumerate(df.iterrows()):
            cells = tbl.rows[r_idx + 1].cells
            cells[0].text = str(idx)
            for j, v in enumerate(row):
                cells[j + 1].text = str(v)
        doc.add_paragraph()

    def add_figure(png: Path, width_in: float = 5.5) -> None:
        if png.exists():
            doc.add_picture(str(png), width=Inches(width_in))
            doc.add_paragraph()

    # ── Cover ──────────────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("Farish Street Cooling Assessment — Year 2")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = font_name if font_name != "DejaVu Sans" else "Calibri"
    doc.add_paragraph("Community Survey Results")
    doc.add_paragraph(f"Generated {dt.date.today().strftime('%B %Y')}")
    doc.add_page_break()

    # ── Sections ───────────────────────────────────────────────────────────────
    sheet_order = ["SocioDem", "Comfortability", "Decision Making", "Green Spaces"]
    first_sheet = True
    for sheet in sheet_order:
        blocks = y2.get(sheet, [])
        if not blocks:
            continue
        heading = SECTION_HEADINGS.get(sheet, sheet)
        if first_sheet:
            add_heading(heading, level=1)
            first_sheet = False
        else:
            add_section_divider(heading)

        for block in blocks:
            qn = block.q_num()
            if is_matrix_question(block.question):
                sub_title = block.subquestion or ""
                add_heading(f"Q{qn} — {sub_title}", level=2)
            else:
                add_heading(block.plot_title().replace("\n", " "), level=2)

            # Count table
            tdf = build_table_df(block)
            add_table(tdf)

            # Chart PNG
            slug = _slug(block.subquestion or block.question)
            # find matching png
            matches = list(OUT.glob(f"y2_{sheet.lower()}*_{slug}*.png"))
            if not matches:
                alt = _slug(block.plot_title())
                matches = list(OUT.glob(f"y2_*{alt[:20]}*.png"))
            if matches:
                add_figure(matches[0])

    doc_path = OUT / "Farish_Survey_Year2_Report.docx"
    doc.save(str(doc_path))
    print(f"Wrote {doc_path}")


# ── Dashboard data builders ───────────────────────────────────────────────────
SHEET_TAB_DEF = [
    ("SocioDem",        "sociodem",        "Socio-Demographics"),
    ("Comfortability",  "comfortability",  "Comfort & Heat Impacts"),
    ("Decision Making", "decision_making", "Decision Making"),
    ("Green Spaces",    "green_spaces",    "Green Space Perceptions"),
]


def build_dashboard_data(
    y1: dict[str, list[QuestionBlock]],
    y2: dict[str, list[QuestionBlock]],
) -> dict:
    result: dict = {}
    for sheet_key, tab_id, tab_label in SHEET_TAB_DEF:
        blocks_y1_map = {b.key(): b for b in y1.get(sheet_key, [])}
        blocks_y2 = y2.get(sheet_key, [])
        questions = []

        # Y2 questions (with optional Y1 match)
        seen_keys: set[str] = set()
        for b2 in blocks_y2:
            b1 = blocks_y1_map.get(b2.key())
            opts2 = b2.option_labels
            n2 = len(opts2)

            def _ser(s: SeriesRow, n: int, opts: list[str]) -> dict:
                pcts = _pct_list(s, n)
                total = sum(c for c in s.counts[:n] if not np.isnan(c))
                return {
                    "label": s.label,
                    "pct":    [round(p, 1) for p in pcts],
                    "counts": [int(round(c)) if not np.isnan(c) else 0 for c in s.counts[:n]],
                    "n":      int(round(total)),
                }

            y2_series = [_ser(s, n2, opts2) for s in b2.series]
            y1_series = []
            opts1 = []
            if b1:
                opts1 = b1.option_labels
                n1 = len(opts1)
                y1_series = [_ser(s, n1, opts1) for s in b1.series]

            questions.append({
                "title":      b2.plot_title().replace("\n", " "),
                "options":    opts2,
                "y1_options": opts1,
                "y2_series":  y2_series,
                "y1_series":  y1_series,
            })
            seen_keys.add(b2.key())

        # Y1-only questions not matched in Y2
        for b1 in y1.get(sheet_key, []):
            if b1.key() in seen_keys:
                continue
            opts1 = b1.option_labels
            n1 = len(opts1)
            y1_series = []
            for s in b1.series:
                pcts = _pct_list(s, n1)
                total = sum(c for c in s.counts[:n1] if not np.isnan(c))
                y1_series.append({
                    "label": s.label,
                    "pct":    [round(p, 1) for p in pcts],
                    "counts": [int(round(c)) if not np.isnan(c) else 0 for c in s.counts[:n1]],
                    "n":      int(round(total)),
                })
            questions.append({
                "title":      b1.plot_title().replace("\n", " "),
                "options":    opts1,
                "y1_options": opts1,
                "y2_series":  [],
                "y1_series":  y1_series,
            })

        result[sheet_key] = {
            "tab_id":    tab_id,
            "label":     tab_label,
            "questions": questions,
        }
    return result


def build_raw_data(
    y1: dict[str, list[QuestionBlock]],
    y2: dict[str, list[QuestionBlock]],
) -> list[dict]:
    rows: list[dict] = []
    for sheet_key, _tab_id, _label in SHEET_TAB_DEF:
        section = SECTION_HEADINGS.get(sheet_key, sheet_key)
        blocks_y1_map = {b.key(): b for b in y1.get(sheet_key, [])}
        for b2 in y2.get(sheet_key, []):
            b1 = blocks_y1_map.get(b2.key())
            opts2 = b2.option_labels
            n2 = len(opts2)
            s2 = next(
                (s for s in b2.series if "all respondents" in s.label.lower()),
                b2.series[0] if b2.series else None,
            )
            s1 = None
            if b1:
                s1 = next(
                    (s for s in b1.series if "all festival" in s.label.lower()),
                    b1.series[-1] if b1.series else None,
                )
            title = b2.plot_title().replace("\n", " ")
            p2_list = _pct_list(s2, n2) if s2 else [None] * n2
            p1_list = _pct_list(s1, len(b1.option_labels)) if (s1 and b1) else [None] * n2
            for j, opt in enumerate(opts2):
                row: dict = {
                    "section":  section,
                    "question": title,
                    "option":   opt,
                }
                if s2 and j < len(s2.counts):
                    v = s2.counts[j]
                    row["y2_count"] = int(round(v)) if not np.isnan(v) else None
                    row["y2_pct"]   = round(p2_list[j], 1) if p2_list[j] is not None else None
                else:
                    row["y2_count"] = row["y2_pct"] = None

                if s1 and b1 and j < len(b1.option_labels) and j < len(s1.counts):
                    v = s1.counts[j]
                    row["y1_count"] = int(round(v)) if not np.isnan(v) else None
                    row["y1_pct"]   = round(p1_list[j], 1) if p1_list[j] is not None else None
                else:
                    row["y1_count"] = row["y1_pct"] = None
                rows.append(row)
    return rows


# ── Dashboard HTML ────────────────────────────────────────────────────────────
def write_dashboard(
    dash_data: dict,
    raw_rows: list[dict],
    meta: dict,
    change_rows: list[dict],
) -> None:
    data_json   = json.dumps(dash_data,   ensure_ascii=False)
    raw_json    = json.dumps(raw_rows,    ensure_ascii=False)
    change_json = json.dumps(change_rows, ensure_ascii=False)
    meta_json   = json.dumps(meta,        ensure_ascii=False, indent=2)

    # Build tab bar HTML
    tab_buttons = ""
    tab_panels  = ""
    for i, (sheet_key, tab_id, tab_label) in enumerate(SHEET_TAB_DEF):
        active_cls = " active" if i == 0 else ""
        tab_buttons += (
            f'<button class="tab-btn{active_cls}" '
            f'data-sheet="{sheet_key}" '
            f'onclick="switchTab(\'{sheet_key}\')">'
            f'{html.escape(tab_label)}</button>\n'
        )
        display = "block" if i == 0 else "none"
        tab_panels += (
            f'<div id="panel-{tab_id}" class="tab-panel" style="display:{display}">'
            f'<!-- charts rendered by JS -->'
            f'</div>\n'
        )
    tab_buttons += '<button class="tab-btn" data-sheet="rawdata" onclick="switchTab(\'rawdata\')">Raw Data</button>\n'
    tab_panels  += '<div id="panel-rawdata" class="tab-panel" style="display:none"></div>\n'

    HTML = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>Farish Street Survey Dashboard</title>
<script src="https://cdn.plot.ly/plotly-latest.min.js"></script>
<style>
*{box-sizing:border-box}
body{font-family:-apple-system,BlinkMacSystemFont,"Segoe UI",Helvetica,Arial,sans-serif;
     margin:0;background:#f4f6f9;color:#1a1a2e;font-size:14px}
header{background:#1f3a5f;color:#fff;padding:20px 28px}
header h1{font-size:1.2rem;margin:0 0 4px;font-weight:700}
header p{margin:0;font-size:.8rem;color:#b8c9dc;line-height:1.5}
.topbar{display:flex;align-items:center;gap:0;background:#fff;
        border-bottom:2px solid #dde3ea;position:sticky;top:0;z-index:100;
        padding:0 20px;flex-wrap:wrap}
.tab-btn{background:none;border:none;border-bottom:3px solid transparent;
         cursor:pointer;font-size:13px;font-weight:600;color:#5c6b7a;
         padding:12px 14px;margin-bottom:-2px;white-space:nowrap;transition:all .15s}
.tab-btn:hover{color:#1f3a5f}
.tab-btn.active{color:#1f3a5f;border-bottom-color:#1f3a5f}
.view-bar{display:flex;gap:6px;margin-left:auto;padding:8px 0}
.vbtn{background:#f1f3f6;border:1px solid #dde3ea;border-radius:20px;
      cursor:pointer;font-size:11px;font-weight:600;color:#5c6b7a;
      padding:5px 12px;transition:all .15s}
.vbtn:hover{background:#e2e8ef}
.vbtn.active{background:#1f3a5f;color:#fff;border-color:#1f3a5f}
main{padding:16px 24px 60px;max-width:1120px;margin:0 auto}
.chart-card{background:#fff;border:1px solid #e2e6ea;border-radius:10px;
            padding:16px 16px 8px;margin-bottom:18px}
.chart-title{font-size:13px;font-weight:700;color:#1a1a2e;margin:0 0 10px;line-height:1.35}
.plotly-chart{width:100%}
.legend-row{display:flex;gap:16px;flex-wrap:wrap;padding:10px 0 6px;align-items:center}
.swatch{display:inline-block;width:13px;height:13px;border-radius:2px;
        margin-right:5px;vertical-align:middle;flex-shrink:0}
.legend-row span{font-size:11.5px;color:#555;display:flex;align-items:center}
table.raw-tbl{width:100%;border-collapse:collapse;font-size:12px}
table.raw-tbl th,table.raw-tbl td{border:1px solid #e2e6ea;padding:6px 8px;text-align:left}
table.raw-tbl th{background:#f1f3f6;font-weight:700;color:#4a5568;position:sticky;top:0}
table.raw-tbl tr:hover td{background:#f9fafb}
td.num{text-align:right;font-variant-numeric:tabular-nums}
.pos{color:#276749}.neg{color:#c53030}
.search-box{width:100%;max-width:360px;padding:8px 12px;border:1px solid #dde3ea;
            border-radius:8px;font-size:13px;margin-bottom:14px;display:block}
pre.meta-pre{font-size:11px;color:#5c6b7a;white-space:pre-wrap;line-height:1.6}
.no-data{color:#999;font-size:12px;padding:20px 0}
</style>
</head>
<body>
<header>
  <h1>Farish Street Community Survey</h1>
  <p>Interactive explorer — Year 1 (2023) &amp; Year 2 (2024/25, n≈162 combined). Use the view toggle to compare years.</p>
</header>
<div class="topbar">
__TAB_BUTTONS__
  <div class="view-bar">
    <button class="vbtn active" id="vbtn-y2"      onclick="setView('y2')">Year 2</button>
    <button class="vbtn"        id="vbtn-y1"      onclick="setView('y1')">Year 1</button>
    <button class="vbtn"        id="vbtn-compare" onclick="setView('compare')">Y1 vs Y2</button>
  </div>
</div>
<main>
  <div class="legend-row">
    <span><span class="swatch" style="background:__C1__"></span>June 22 (Y1)</span>
    <span><span class="swatch" style="background:__C2__"></span>August 25 (Y1)</span>
    <span><span class="swatch" style="background:__C3__"></span>All festival dates (Y1)</span>
    <span><span class="swatch" style="background:__CY2__"></span>Year 2 (all respondents)</span>
    <span><span class="swatch" style="background:__CY2C__"></span>Year 2 in comparisons</span>
  </div>
__TAB_PANELS__
</main>
<script>
const DATA = __DATA_JSON__;
const RAW  = __RAW_JSON__;
const CHANGE = __CHANGE_JSON__;

const C1='__C1__', C2='__C2__', C3='__C3__', CY2='__CY2__', CY2C='__CY2C__';
const SHEET_IDS = {
  "SocioDem":"sociodem","Comfortability":"comfortability",
  "Decision Making":"decision_making","Green Spaces":"green_spaces"
};
const SHEETS = Object.keys(SHEET_IDS);

let activeSheet = "SocioDem";
let viewMode = "y2";

function switchTab(sheetKey){
  SHEETS.forEach(s=>{
    const el=document.getElementById('panel-'+SHEET_IDS[s]);
    if(el) el.style.display='none';
  });
  const rawEl=document.getElementById('panel-rawdata');
  if(rawEl) rawEl.style.display='none';
  document.querySelectorAll('.tab-btn').forEach(b=>b.classList.remove('active'));
  const activeBtn=document.querySelector('[data-sheet="'+sheetKey+'"]');
  if(activeBtn) activeBtn.classList.add('active');

  if(sheetKey==='rawdata'){
    if(rawEl) rawEl.style.display='block';
    renderRawData();
  } else {
    const panelId='panel-'+SHEET_IDS[sheetKey];
    const p=document.getElementById(panelId);
    if(p) p.style.display='block';
    activeSheet=sheetKey;
    renderTabCharts(sheetKey);
  }
}

function setView(mode){
  viewMode=mode;
  ['y2','y1','compare'].forEach(m=>{
    const b=document.getElementById('vbtn-'+m);
    if(b) b.classList.toggle('active', m===mode);
  });
  if(activeSheet!=='rawdata') renderTabCharts(activeSheet);
}

function renderTabCharts(sheetKey){
  const sheet=DATA[sheetKey];
  if(!sheet) return;
  const panelId='panel-'+SHEET_IDS[sheetKey];
  const panel=document.getElementById(panelId);
  if(!panel) return;

  // Build cards if not already present
  if(!panel.dataset.built){
    panel.dataset.built='1';
    let html='';
    sheet.questions.forEach(function(q,qi){
      const cid='chart-'+SHEET_IDS[sheetKey]+'-'+qi;
      html+='<div class="chart-card">'+
            '<p class="chart-title">'+esc(q.title)+'</p>'+
            '<div id="'+cid+'" class="plotly-chart"></div>'+
            '</div>';
    });
    panel.innerHTML=html;
  }

  sheet.questions.forEach(function(q,qi){
    const cid='chart-'+SHEET_IDS[sheetKey]+'-'+qi;
    const el=document.getElementById(cid);
    if(!el) return;
    const traces=buildTraces(q);
    if(!traces.length){
      el.innerHTML='<p class="no-data">No data available for this view.</p>';
      return;
    }
    const nOpts=Math.max(
      traces[0].x ? traces[0].x.length : 0, 1
    );
    const layout={
      barmode:'group',
      xaxis:{
        title:{text:'Response',font:{size:11},standoff:8},
        tickfont:{size:10},automargin:true,
        tickangle: nOpts>5 ? -30 : 0
      },
      yaxis:{
        title:{text:'Percent of respondents',font:{size:11}},
        ticksuffix:'%',rangemode:'tozero',
        gridcolor:'#e5e5e5',gridwidth:1,
        zeroline:false
      },
      legend:{orientation:'h',y:-0.28,font:{size:11},xanchor:'center',x:0.5},
      margin:{t:10,b:90,l:60,r:20},
      paper_bgcolor:'white',plot_bgcolor:'white',
      height:330,
      showlegend:traces.length>1
    };
    Plotly.react(cid, traces, layout, {responsive:true,displayModeBar:false});
  });
}

function buildTraces(q){
  const traces=[];
  if(viewMode==='y2'){
    q.y2_series.forEach(function(s){
      const total=s.n>0?s.n:s.counts.reduce(function(a,b){return a+b;},0);
      traces.push({
        type:'bar', name:s.label+(total>0?' (n='+total+')':''),
        x:q.options, y:s.pct,
        marker:{color:CY2},
        customdata:s.counts,
        hovertemplate:'<b>%{x}</b><br>%{y:.1f}%  (n=%{customdata})<extra></extra>'
      });
    });
  } else if(viewMode==='y1'){
    const cols=[C1,C2,C3];
    const opts=q.y1_options&&q.y1_options.length?q.y1_options:q.options;
    q.y1_series.forEach(function(s,si){
      const total=s.n>0?s.n:s.counts.reduce(function(a,b){return a+b;},0);
      traces.push({
        type:'bar', name:s.label+(total>0?' (n='+total+')':''),
        x:opts, y:s.pct,
        marker:{color:cols[si%3]},
        customdata:s.counts,
        hovertemplate:'<b>%{x}</b><br>%{y:.1f}%  (n=%{customdata})<extra></extra>'
      });
    });
  } else {
    // Compare: Y1 all festival dates vs Y2 aggregate
    const s1=q.y1_series.find(function(s){return s.label.toLowerCase().indexOf('all festival')>=0;})
           ||(q.y1_series.length?q.y1_series[q.y1_series.length-1]:null);
    const s2=q.y2_series.find(function(s){return s.label.toLowerCase().indexOf('all respondents')>=0;})
           ||(q.y2_series.length?q.y2_series[0]:null);
    const opts1=q.y1_options&&q.y1_options.length?q.y1_options:q.options;
    if(s1){
      traces.push({
        type:'bar', name:'Year 1 — all festival dates',
        x:opts1, y:s1.pct, marker:{color:C3},
        customdata:s1.counts,
        hovertemplate:'<b>%{x}</b><br>Year 1: %{y:.1f}%  (n=%{customdata})<extra></extra>'
      });
    }
    if(s2){
      traces.push({
        type:'bar', name:'Year 2 — all respondents',
        x:q.options, y:s2.pct, marker:{color:CY2C},
        customdata:s2.counts,
        hovertemplate:'<b>%{x}</b><br>Year 2: %{y:.1f}%  (n=%{customdata})<extra></extra>'
      });
    }
  }
  return traces;
}

function renderRawData(){
  const panel=document.getElementById('panel-rawdata');
  if(panel.dataset.built) return;
  panel.dataset.built='1';

  let h='<input class="search-box" type="text" id="raw-search" placeholder="Search questions, options…" oninput="filterRaw()">';
  h+='<div style="overflow-x:auto"><table class="raw-tbl" id="raw-tbl">';
  h+='<thead><tr><th>Section</th><th>Question</th><th>Option</th>';
  h+='<th class="num">Y1 count<br>(all dates)</th><th class="num">Y1 %</th>';
  h+='<th class="num">Y2 count</th><th class="num">Y2 %</th>';
  h+='<th class="num">Δ pp</th></tr></thead><tbody>';
  RAW.forEach(function(r){
    const d=(r.y1_pct!==null&&r.y2_pct!==null)?Math.round((r.y2_pct-r.y1_pct)*10)/10:null;
    const cls=d===null?'':d>0.5?'pos':d<-0.5?'neg':'';
    h+='<tr>'+
       '<td>'+esc(r.section)+'</td>'+
       '<td>'+esc(r.question)+'</td>'+
       '<td>'+esc(r.option)+'</td>'+
       '<td class="num">'+(r.y1_count!==null?r.y1_count:'—')+'</td>'+
       '<td class="num">'+(r.y1_pct!==null?r.y1_pct+'%':'—')+'</td>'+
       '<td class="num">'+(r.y2_count!==null?r.y2_count:'—')+'</td>'+
       '<td class="num">'+(r.y2_pct!==null?r.y2_pct+'%':'—')+'</td>'+
       '<td class="num '+cls+'">'+(d!==null?(d>0?'+':'')+d:'—')+'</td>'+
       '</tr>';
  });
  h+='</tbody></table></div>';
  panel.innerHTML=h;
}

function filterRaw(){
  const q=document.getElementById('raw-search').value.toLowerCase();
  const rows=document.querySelectorAll('#raw-tbl tbody tr');
  rows.forEach(function(r){
    r.style.display=r.textContent.toLowerCase().indexOf(q)>=0?'':'none';
  });
}

function esc(s){
  return String(s).replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;');
}

// Initial render
document.addEventListener('DOMContentLoaded',function(){
  renderTabCharts(activeSheet);
});
</script>
</body>
</html>"""

    # Inject values
    HTML = (HTML
        .replace("__TAB_BUTTONS__", tab_buttons)
        .replace("__TAB_PANELS__",  tab_panels)
        .replace("__DATA_JSON__",   data_json)
        .replace("__RAW_JSON__",    raw_json)
        .replace("__CHANGE_JSON__", change_json)
        .replace("__C1__",   COL_DATE1,  3)
        .replace("__C2__",   COL_DATE2,  3)
        .replace("__C3__",   COL_ALL,    3)
        .replace("__CY2__",  COL_Y2,     3)
        .replace("__CY2C__", COL_Y2_COMP, 3)
        .replace("__META_JSON__", html.escape(meta_json))
    )
    (ROOT / "survey_dashboard.html").write_text(HTML, encoding="utf-8")


# ── Main ──────────────────────────────────────────────────────────────────────
def main() -> None:
    font = setup_fonts()
    apply_style(font)
    print(f"Font: {font!r}")

    OUT.mkdir(parents=True, exist_ok=True)
    for old in OUT.glob("*.png"):
        old.unlink()

    y1 = load_all_blocks(Y1_PATH, "Year 1")
    y2 = load_all_blocks(Y2_PATH, "Year 2")
    comp = build_comparison_index(y1, y2)
    change_rows = build_change_summaries(comp)
    (OUT / "survey_comparison_rows.json").write_text(
        json.dumps(change_rows, indent=2, ensure_ascii=False), encoding="utf-8"
    )

    # ── Year 2 charts ──────────────────────────────────────────────────────────
    rel_y2: list[str] = []
    png_map: dict[str, Path] = {}
    for sheet, blocks in y2.items():
        for bi, block in enumerate(blocks):
            slug = _slug(block.subquestion or block.plot_title())
            fname = f"y2_{sheet.lower()}_{bi:02d}_{slug}.png"
            p = OUT / fname
            plot_block(block, p)
            rel_y2.append(f"Outputs/{fname}")
            png_map[block.key()] = p
            print(f"  {fname}")

    # ── Year 1 (individual, matching Year 2 style) ────────────────────────────
    for sheet, blocks in y1.items():
        for bi, block in enumerate(blocks):
            slug = _slug(block.subquestion or block.plot_title())
            fname = f"y1_{sheet.lower()}_{bi:02d}_{slug}.png"
            plot_block(block, OUT / fname, year_label="Year 1")

    # ── Comparison charts ─────────────────────────────────────────────────────
    rel_comp: list[str] = []
    for idx, (sheet, key, b1, b2) in enumerate(comp):
        if "10. if yes" in b1.question.lower():
            continue
        try:
            fname = f"compare_{idx:02d}_{_slug(key)}.png"
            plot_y1_y2_comparison(b1, b2, OUT / fname)
            rel_comp.append(f"Outputs/{fname}")
        except Exception as e:
            print(f"  skip compare {idx}: {e}")

    meta = {
        "y1_file": Y1_PATH.name,
        "y2_file": Y2_PATH.name,
        "note": (
            "Year 2 data contains a single pooled sample (n≈162). "
            "No event-date breakdown is available for Year 2, which is why "
            "socio-demographic charts show one bar instead of multiple dates."
        ),
        "font": font,
        "y2_figures": len(rel_y2),
        "comparison_figures": len(rel_comp),
        "comparison_table_rows": sum(len(r["options"]) for r in change_rows),
        "colors": {
            "June 22":            COL_DATE1,
            "August 25":          COL_DATE2,
            "All festival dates": COL_ALL,
            "Year 2 (single)":    COL_Y2,
            "Year 2 (vs Y1)":     COL_Y2_COMP,
        },
    }

    dash_data = build_dashboard_data(y1, y2)
    raw_rows  = build_raw_data(y1, y2)
    write_dashboard(dash_data, raw_rows, meta, change_rows)
    print(f"\nWrote dashboard → {ROOT/'survey_dashboard.html'}")

    build_word_doc(y2, png_map, font)


if __name__ == "__main__":
    main()
